#coding=utf-8
import os
from docx import Document
from docx.shared import Inches
import sys
from IgnoreArray import IgnoreArray
from ExtDetector import is_img

lines=0
ignoreArr = IgnoreArray()
#print ignoreArr
def eachFile(filePath, document):
		global lines
		if (os.path.isfile(filePath)):
			(path, name) = os.path.split(filePath)
			document.add_heading(name, 1)
			if is_img(filePath):
				document.add_picture(filePath, width=Inches(1.25))
				return
			p=document.add_paragraph('')
			fopen =open(filePath, 'r')
			for line in fopen:
				if(line.strip() == ''): continue
				if(line.strip() == '\r'): continue
				if(line.strip() == '\n'): continue
				if(line.strip() == '\r\n'): continue
				if(line.strip() == '\u0x0D'): continue
				if(line.strip() == '\u0x0A'): continue
				lines=lines+1
				#print line
				try:
					p.add_run('%s\r\n'%(line.decode('utf-8')))
				except BaseException:
					print 'error'
			print name
		else:
			files = os.listdir(filePath)
			for f in files:
				if(f in ignoreArr): continue
				eachFile(filePath + '/' + f, document)
document = Document()

document.add_heading('source code'.decode('utf-8'), 0)
if(len(sys.argv) < 3):
	print 'Without src path and output path'
else:
	if(os.path.exists(sys.argv[1])):
		if(not os.path.exists(sys.argv[2])):
			eachFile(sys.argv[1], document)
			document.save(sys.argv[2])

			print 'Merge all the files in ' + sys.argv[1] + ' to ' + sys.argv[2] + ' success'
			print 'All source are ' + str(lines) + ' lines'

		else:
			print sys.argv[2] + 'has exist'
	else:
		print sys.argv[1] + ' not exist'








